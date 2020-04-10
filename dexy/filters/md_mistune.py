from dexy.filter import DexyFilter
import re

try:
    import mistune
    MISTUNE_AVAILABLE = True
except ImportError:
    MISTUNE_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.shared import Cm
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class MistuneFilter(DexyFilter):
    """
    Runs mistune processor to convert markdown to HTML.

    """
    aliases = ['mistune']
    _settings = {
            'input-extensions' : ['.md', '.txt'],
            'output-extensions' : ['.html'],
            'escape' : True,
            'hard_wrap' : False,
            'use_xhtml' : False,
            'parse_block_html' : False,
            'parse_inline_html' : False
            }

    def is_active(self):
        return MISTUNE_AVAILABLE

    def process_text(self, input_text):
        return mistune.markdown(
                input_text,
                escape = self.setting('escape'),
                hard_wrap = self.setting('hard_wrap'),
                use_xhtml = self.setting('use_xhtml'),
                parse_block_html = self.setting('parse_block_html'),
                parse_inline_html = self.setting('parse_inline_html')
                )


class MistuneDocxFilter(DexyFilter):
    """
    Runs mistune processor to convert markdown to HTML.

    """
    aliases = ['md2docx']
    _settings = {
            'input-extensions' : ['.md'],
            'output-extensions' : ['.docx'],
            'output' : True,
            'escape' : True,
            'hard_wrap' : False,
            'use_xhtml' : False,
            'parse_block_html' : False,
            'parse_inline_html' : False
            }

    def is_active(self):
        return MISTUNE_AVAILABLE and DOCX_AVAILABLE

    def process(self):
        input_text = str(self.input_data)
        renderer = PythonDocxRenderer(Document())
        mistune.markdown(input_text, renderer = renderer)
        renderer.document.save(self.output_data.storage.data_file())

class MathBlockGrammar(mistune.BlockGrammar):
    block_math = re.compile(r"^\$\$(.*?)\$\$", re.DOTALL)


class MathBlockLexer(mistune.BlockLexer):
    default_rules = ['block_math'] + mistune.BlockLexer.default_rules

    def __init__(self, rules=None, **kwargs):
        if rules is None:
            rules = MathBlockGrammar()
        super(MathBlockLexer, self).__init__(rules, **kwargs)

    def parse_block_math(self, m):
        """Parse a $$math$$ block"""
        self.tokens.append({'type': 'block_math', 'text': m.group(1)})


class MarkdownWithMath(mistune.Markdown):
    def __init__(self, renderer, **kwargs):
        kwargs['block'] = MathBlockLexer
        super(MarkdownWithMath, self).__init__(renderer, **kwargs)

    def output_block_math(self):
        return self.renderer.block_math(self.token['text'])

class TextSpan(object):
    def __init__(self, text):
        self.text = text

# Inpsired by https://github.com/mjanv/mistune-docx
class PythonDocxRenderer(mistune.Renderer):
    def __init__(self, document, **kwargs):
        super(PythonDocxRenderer, self).__init__(**kwargs)
        self.table_memory = []
        self.img_counter = 0
        self.document = document
        self.buffer = []
        self.list_buffer = []

    def header(self, text, level, raw):
        self.document.add_heading(" ".join(b.get('text') for b in self.buffer), level=(level - 1))
        self.buffer = []
        return ""

    def paragraph(self, text, buf=None, style=None, do_return=False):
        if style:
            p = self.document.add_paragraph(style=style)
        else:
            p = self.document.add_paragraph()

        for d in (self.buffer or buf or []):
            run = p.add_run(d['text'])
            run.italic = d.get('italic', False)
            run.bold = d.get('bold', False)

        self.buffer = []

        if do_return:
            return p
        else:
            return ""

    def list(self, body, ordered):
        if ordered:
            list_style = "List Number"
        else:
            list_style = "List Bullet"

        for buf in self.list_buffer:
            p = self.paragraph("", buf=buf, style=list_style, do_return=True)

        self.list_buffer = []

        run = p.add_run()
        run.add_break()
        return ""

    def list_item(self, text):
        self.list_buffer.append(self.buffer)
        self.buffer = []
        return ""
#
#    def table(self, header, body):
#        number_cols = header.count('\n') - 2
#        number_rows = int(len(self.table_memory) / number_cols)
#        cells = ["table.rows[%d].cells[%d].paragraphs[0]%s\n" % (i, j, self.table_memory.pop(0)[1:]) for i, j in itertools.product(range(number_rows), range(number_cols))]
#        return '\n'.join(["table = document.add_table(rows=%d, cols=%d, style = 'BasicUserTable')" % (number_rows, number_cols)] + cells) + 'document.add_paragraph().add_run().add_break()\n'
#
#    def table_cell(self, content, **flags):
#        self.table_memory.append(content)
#        return content
#
#    # SPAN LEVEL
    def text(self, text):
        self.buffer.append({'text' : text})
        return ""

    def emphasis(self, text):
        self.buffer[-1]['italic'] = True
        return ""

    def double_emphasis(self, text):
        self.buffer[-1]['bold'] = True
        return ""

#    def block_code(self, code, language):
#        code = code.replace('\n', '\\n')
#        return "p = document.add_paragraph()\np.add_run(\"%s\")\np.style = 'BasicUserQuote'\np.add_run().add_break()\n" % code
#
    def link(self, link, title, content):
        self.buffer.append({"text" : " (%s)" % (link)})
        self.paragraph("")
        return ""

    def image(self, src, title, alt_text):
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_after = Pt(18)

        run = p.add_run()
        run.add_picture(src, width=Cm(15))
        run.add_break()
        self.buffer.append({'text' : alt_text, 'italic' : True})
        return ""

    def hrule(self):
        self.document.add_page_break()
        return ""

#    def block_math(self, text):
#        import sympy
#        if not os.path.exists('tmp'):
#            os.makedirs('tmp')
#        filename = 'tmp/tmp%d.png' % self.img_counter
#        self.img_counter = self.img_counter + 1
#        sympy.preview(r'$$%s$$' % text, output='png', viewer='file', filename=filename, euler=False)
#        return self.image(filename, None, "Equation " + str(self.img_counter - 1))
